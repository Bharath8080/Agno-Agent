import streamlit as st 
from phi.agent import Agent
from phi.model.google import Gemini
from linkup import LinkupClient
from google.generativeai import upload_file, get_file
import google.generativeai as genai
import time
import os
import tempfile
from pathlib import Path
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing

# Load environment variables
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
LINKUP_API_KEY = os.getenv("LINKUP_API_KEY")
if API_KEY:
    genai.configure(api_key=API_KEY)

# Initialize Linkup client
linkup_client = LinkupClient(api_key=LINKUP_API_KEY)

# Page configuration
st.set_page_config(
    page_title="Multimodal_AI_Agent",
    page_icon="https://framerusercontent.com/images/dSEDRe8n4rNO7Cmt0WaTx4e0iyQ.jpeg",
    layout="wide"
)

st.markdown(
    """
    <h2 style="display: flex; align-items: center;">
        <img src="https://agno-public.s3.us-east-1.amazonaws.com/assets/logo-dark.svg" width="125" style="margin-right: 7px;"> 
        Powered by 
        <img src="https://logospng.org/download/google-gemini/google-gemini-1024.png" width="125" style="margin-left: 2px;"><img src="https://camo.githubusercontent.com/77ba4ba362fc39151379e4e7691125c8bb130eb2ade811ce9f76d4d5236c6847/68747470733a2f2f75706c6f61642e77696b696d656469612e6f72672f77696b6970656469612f636f6d6d6f6e732f7468756d622f662f66302f476f6f676c655f426172645f6c6f676f2e7376672f3132303070782d476f6f676c655f426172645f6c6f676f2e7376672e706e67" width="30" style="margin-left: 0px;">
        
    </h2>
    """,
    unsafe_allow_html=True
)


# Define a search function that uses the Linkup client directly
def perform_linkup_search(query, depth="standard", output_type="sourcedAnswer"):
    try:
        response = linkup_client.search(
            query=query,
            depth=depth,
            output_type=output_type,
        )
        return response
    except Exception as e:
        st.error(f"Linkup search error: {e}")
        return {"error": str(e)}

@st.cache_resource
def initialize_agent():
    return Agent(
        name="Multimodal AI Summarizer & Chat",
        model=Gemini(id="gemini-2.0-flash-exp"),
        markdown=True,
    )

# Initialize the agent
multimodal_Agent = initialize_agent()

# Initialize a specific agent for document analysis with lower temperature
@st.cache_resource
def initialize_document_agent():
    return Agent(
        name="Document Analysis Agent",
        model=Gemini(
            id="gemini-2.0-flash-exp",
            temperature=0.1  # Low temperature for more precise responses
        ),
        markdown=True,
    )

# Initialize a basic chat agent for general queries
@st.cache_resource
def initialize_chat_agent():
    return Agent(
        name="General Chat Agent",
        model=Gemini(
            id="gemini-2.0-flash-exp",
            temperature=0.7  # Balanced temperature for general conversations
        ),
        markdown=True,
    )

# Initialize the document agent
document_Agent = initialize_document_agent()

# Initialize the chat agent
chat_Agent = initialize_chat_agent()

# Initialize session state variables
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'chat_memory' not in st.session_state:
    st.session_state.chat_memory = []  # Separate memory for agent context
if 'processed_file' not in st.session_state:
    st.session_state.processed_file = None
if 'document_text' not in st.session_state:
    st.session_state.document_text = None
if 'file_type' not in st.session_state:
    st.session_state.file_type = None
if 'mode' not in st.session_state:
    st.session_state.mode = "Direct Chat"  # Default mode is chat

# Set a fixed memory length instead of using a slider
memory_length = 5

# Function to format chat history for agent context
def format_chat_memory_for_agent(chat_memory, max_turns=memory_length):
    formatted_history = ""
    # Take only the last few turns to avoid context length issues
    recent_memory = chat_memory[-max_turns*2:] if len(chat_memory) > max_turns*2 else chat_memory
    
    for message in recent_memory:
        role = "USER" if message["role"] == "user" else "ASSISTANT"
        formatted_history += f"{role}: {message['content']}\n\n"
    
    return formatted_history

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

# Function to extract text from TXT
def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if utf-8 fails
        try:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error reading text file: {e}")
            return ""
    except Exception as e:
        st.error(f"Error reading text file: {e}")
        return ""

# Generate a preview for documents
def get_document_preview(file_path, file_extension):
    if file_extension == "pdf":
        try:
            doc = fitz.open(file_path)
            first_page = doc[0]
            pix = first_page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))  # Reduce size for preview
            img_bytes = pix.tobytes("png")
            doc.close()
            return img_bytes, "image"
        except:
            return "PDF Preview Unavailable", "text"
    elif file_extension == "docx":
        try:
            doc = docx.Document(file_path)
            preview_text = ""
            # Get first 5 paragraphs or fewer if document is smaller
            for i, para in enumerate(doc.paragraphs):
                if i < 5 and para.text.strip():
                    preview_text += para.text + "\n\n"
                if i >= 5:
                    preview_text += "..."
                    break
            return preview_text or "Document is empty or contains non-text content", "text"
        except:
            return "DOCX Preview Unavailable", "text"
    elif file_extension == "txt":
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                lines = file.readlines()
                # Get first 10 lines or fewer if file is smaller
                preview_lines = lines[:10]
                preview_text = "".join(preview_lines)
                if len(lines) > 10:
                    preview_text += "..."
                return preview_text or "Text file is empty", "text"
        except:
            return "TXT Preview Unavailable", "text"
    else:
        return "Preview not available for this file type", "text"

# Sidebar logo
st.sidebar.image("https://framerusercontent.com/images/wLLGrlJoyqYr9WvgZwzlw91A8U.png", use_container_width=True)

# Unified file selection section
st.sidebar.header("📤 Upload a File or Chat Directly")
file_type_options = ["💬 GeneralChat", "🖼️ Image/Video", "📄 Documents (PDF, TXT, DOCX)"]
file_type = st.sidebar.radio("Select file type", file_type_options)

if file_type == "💬 GeneralChat":
    st.session_state.file_type = "direct_chat"
    st.session_state.processed_file = None
    st.session_state.document_text = None
    st.session_state.mode = "Direct Chat"
    
elif file_type == "🖼️ Image/Video":
    st.session_state.file_type = "media"
    st.session_state.mode = "File Analysis"
    
    # Show file uploader for Image/Video
    file = st.sidebar.file_uploader(
        "Upload an image or video file", type=['jpg', 'jpeg', 'png', 'mp4', 'mov', 'avi']
    )
   
    if file:
        # If file type changed, clear chat history
        if st.session_state.get('last_file_name') != file.name:
            st.session_state.chat_history = []
            st.session_state.chat_memory = []
            st.session_state['last_file_name'] = file.name
            
        file_extension = file.name.split(".")[-1].lower()
        temp_suffix = f".{file_extension}"
       
        with tempfile.NamedTemporaryFile(delete=False, suffix=temp_suffix) as temp_file:
            temp_file.write(file.read())
            file_path = temp_file.name
       
        if file_extension in ['jpg', 'jpeg', 'png']:
            st.sidebar.image(file_path, caption="Uploaded Image", use_container_width=True)
        elif file_extension in ['mp4', 'mov', 'avi']:
            st.sidebar.video(file_path, format="video/mp4", start_time=0)
       
        st.session_state.processed_file = file_path
        st.sidebar.success("Media file uploaded successfully!")
       
elif file_type == "📄 Documents (PDF, TXT, DOCX)":
    st.session_state.file_type = "document"
    st.session_state.mode = "File Analysis"
    
    # Show file uploader for PDF, TXT and DOCX
    file = st.sidebar.file_uploader("Upload a document file", type=['pdf', 'txt', 'docx'])
   
    if file:
        # If file changed, clear chat history
        if st.session_state.get('last_file_name') != file.name:
            st.session_state.chat_history = []
            st.session_state.chat_memory = []
            st.session_state['last_file_name'] = file.name
            
        file_extension = file.name.split(".")[-1].lower()
        temp_suffix = f".{file_extension}"
       
        with tempfile.NamedTemporaryFile(delete=False, suffix=temp_suffix) as temp_file:
            temp_file.write(file.read())
            file_path = temp_file.name
       
        with st.spinner(f"Processing {file_extension.upper()} file..."):
            # Extract text based on file type
            if file_extension == "pdf":
                document_text = extract_text_from_pdf(file_path)
            elif file_extension == "docx":
                document_text = extract_text_from_docx(file_path)
            elif file_extension == "txt":
                document_text = extract_text_from_txt(file_path)
            else:
                document_text = "Unsupported file format"
           
            # Store in session state
            st.session_state.document_text = document_text
            st.session_state.processed_file = file_path
           
            # Add document preview
            preview_content, preview_type = get_document_preview(file_path, file_extension)
            if preview_type == "image":
                st.sidebar.image(preview_content, caption=f"{file_extension.upper()} Preview", use_container_width=True)
            else:
                st.sidebar.text_area("Document Preview", preview_content, height=200, disabled=True)
           
            # Display document info
            text_length = len(document_text)
            st.sidebar.info(f"{file_extension.upper()} processed: {file.name} ({text_length} characters)")
           
            st.sidebar.success(f"{file_extension.upper()} file processed successfully!")

# Sidebar controls
st.sidebar.header("⚙ Controls")
if st.sidebar.button('🧹 Clear Chat History'):
    st.session_state.chat_history = []
    st.session_state.chat_memory = []
    st.rerun()
if st.sidebar.button('🗑 Clear All Data'):
    st.session_state.processed_file = None
    st.session_state.chat_history = []
    st.session_state.chat_memory = []
    st.session_state.document_text = None
    st.session_state.file_type = None
    st.session_state.mode = "Direct Chat"
    if 'last_file_name' in st.session_state:
        del st.session_state['last_file_name']
    st.sidebar.success("Data cleared!")
    st.rerun()

# Display chat history
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Chat input
input_placeholder = "Ask a question..." if st.session_state.file_type == "direct_chat" else "Ask a question about your file..."
if query := st.chat_input(input_placeholder):
    st.session_state.chat_history.append({"role": "user", "content": query})
    st.session_state.chat_memory.append({"role": "user", "content": query})
    with st.chat_message("user"):
        st.markdown(query)
    
    with st.chat_message("assistant"):
        try:
            # Get formatted chat history for context
            chat_context = format_chat_memory_for_agent(st.session_state.chat_memory, max_turns=memory_length)
            
            # Handle direct chat mode
            if st.session_state.mode == "Direct Chat":
                with st.spinner("Processing your question..."):
                    # First, have the agent assess whether it needs web search
                    assessment_prompt = f"""
                    CHAT HISTORY:
                    {chat_context}
                    
                    USER QUERY: {query}
                    
                    First, assess if this query requires up-to-date information, specific factual details, or 
                    information that might be beyond your knowledge cutoff.
                    
                    Consider these factors:
                    - Is the query about very recent events, current data, or trending topics?
                    - Does it require specific factual information you might not have?
                    - Would searching for information significantly improve the accuracy of your response?
                    - Does it reference previous parts of our conversation?
                    
                    RESPOND ONLY WITH:
                    "SEARCH_NEEDED" - If web search would significantly improve the response quality
                    "DIRECT_ANSWER" - If you can confidently answer without additional information
                    """
                    
                    search_decision = chat_Agent.run(assessment_prompt)
                    
                    # Based on the assessment, either perform search or answer directly
                    if "SEARCH_NEEDED" in search_decision.content:
                        with st.spinner("Searching for information..."):
                            search_results = perform_linkup_search(query, depth="deep")
                            
                            # Create prompt with search results and chat history
                            search_prompt = f"""
                            CHAT HISTORY:
                            {chat_context}
                            
                            USER QUERY: {query}
                            
                            Based on the following search results:
                            
                            {search_results}
                            
                            Please provide a comprehensive answer to the query. Remember to consider our previous conversation
                            when relevant. Format your response clearly and cite information from these search results when appropriate.
                            """
                            
                            # Get response with search results
                            final_response = chat_Agent.run(search_prompt)
                            answer = final_response.content
                    else:
                        # Answer directly from model knowledge
                        direct_prompt = f"""
                        CHAT HISTORY:
                        {chat_context}
                        
                        USER QUERY: {query}
                        
                        Provide a helpful, accurate response based on your knowledge and our previous conversation.
                        If you're uncertain about any aspect, mention this uncertainty.
                        """
                        
                        response = chat_Agent.run(direct_prompt)
                        answer = response.content
            
            # Handle file analysis mode
            else:
                file_path = st.session_state.get('processed_file')
                file_type = st.session_state.get('file_type')
                
                if not file_path:
                    answer = "No file uploaded. Please upload a file first or switch to Direct Chat mode."
                elif file_type == "document":
                    # Document processing with custom prompt and low temperature
                    with st.spinner("Processing your question about the document..."):
                        document_text = st.session_state.document_text
                        file_extension = file_path.split(".")[-1].lower()
                        
                        # If text is too long, truncate it
                        max_text_length = 30000
                        if len(document_text) > max_text_length:
                            document_text = document_text[:max_text_length] + "... [text truncated due to length]"
                        
                        # First check if we need to search for information not in the document
                        need_web_search = True
                        
                        # Create the base prompt for document analysis
                        base_prompt = f"""
                        CHAT HISTORY:
                        {chat_context}
                        
                        You are a precise document analysis system specialized in extracting accurate information from {file_extension.upper()} documents.
                        
                        DOCUMENT CONTENT:
                        {document_text}
                        
                        USER QUERY: {query}
                        
                        TASK:
                        1. Consider our previous conversation for context when answering this query.
                        2. Analyze the content carefully and focus only on information that is explicitly present in the document.
                        3. If the exact answer is present in the document, provide it with direct quotes if appropriate.
                        4. If the answer isn't explicitly stated but can be reasonably inferred from the content, explain your reasoning.
                        """
                        
                        # First check if the answer is in the document
                        initial_check_prompt = base_prompt + """
                        5. IMPORTANT: Only answer "YES" if the answer to the query can be found in the document content.
                        Answer "NO" if the information needed to answer the query is not in the document.
                        """
                        
                        # Check if we need to search web
                        initial_response = document_Agent.run(initial_check_prompt)
                        if "YES" in initial_response.content.upper():
                            need_web_search = False
                        
                        # If answer not in document, perform web search
                        web_search_results = ""
                        if need_web_search:
                            search_query = f"Information about: {query}"
                            with st.spinner("Searching for additional information..."):
                                search_results = perform_linkup_search(search_query, depth="deep", output_type="searchResults")
                                if "error" not in search_results:
                                    web_search_results = "\n\nWEB SEARCH RESULTS:\n" + str(search_results)
                        
                        # Now create the full prompt with search results if needed
                        full_prompt = base_prompt + """
                        5. If the answer cannot be found in the document at all, DO NOT attempt to guess.
                        """
                        
                        if web_search_results:
                            full_prompt += f"""
                            6. Use the following web search results to supplement information not found in the document:
                            {web_search_results}
                            
                            When using information from web search, clearly indicate: "This information was not found in the document. Based on web search:"
                            """
                        
                        full_prompt += """
                        FORMAT YOUR RESPONSE:
                        - Start with a direct answer to the query if available
                        - Support with relevant quotes or references from the document
                        - Be factual and precise
                        - If using web search, clearly separate document-based information from web-sourced information
                        - Maintain continuity with our previous conversation when appropriate
                        """
                        
                        # Use the document-specific agent with temperature 0.1
                        response = document_Agent.run(full_prompt)
                        answer = response.content
                else:  # Media files
                    file_extension = file_path.split(".")[-1].lower()
                    
                    # First determine if the query is related to the media or not
                    with st.spinner("Understanding your query..."):
                        # Define a classifier prompt to determine if the query is about the uploaded media
                        classifier_prompt = f"""
                        CHAT HISTORY:
                        {chat_context}
                        
                        USER QUERY: {query}
                        
                        TASK: Determine if the query is specifically asking about the content, context, elements, or analysis 
                        of the uploaded {'image' if file_extension in ['jpg', 'jpeg', 'png'] else 'video'}.
                        Consider our previous conversation and determine if this query is a follow-up about the media.
                        
                        RESPOND ONLY WITH:
                        "MEDIA_RELATED" - If the query is directly asking about the uploaded media content, analysis, or elements, or is a follow-up to previous media analysis
                        "GENERAL_QUESTION" - If the query is a general question not specifically related to analyzing the media content
                        """
                        
                        # Run the classifier prompt to determine query type
                        query_type_response = multimodal_Agent.run(classifier_prompt)
                        
                        # Check if query is media related or general
                        if "MEDIA_RELATED" in query_type_response.content:
                            # This is a media-related query, so we'll analyze the media
                            with st.spinner("Analyzing media for your query..."):
                                processed_file = upload_file(file_path)
                                while processed_file.state.name == "PROCESSING":
                                    time.sleep(1)
                                    processed_file = get_file(processed_file.name)
                                
                                # Create prompt for media analysis
                                media_prompt = f"""
                                CHAT HISTORY:
                                {chat_context}
                                
                                Analyze the uploaded {'image' if file_extension in ['jpg', 'jpeg', 'png'] else 'video'}.
                                
                                USER QUERY: {query}
                                
                                Provide a detailed response that directly addresses the query about this media.
                                Consider our previous conversation for context when giving your response.
                                If your analysis cannot fully answer the query, indicate what additional information 
                                might be needed to provide a complete answer.
                                """
                                
                                # Get the response for media analysis
                                response = multimodal_Agent.run(
                                    media_prompt, 
                                    images=[processed_file] if file_extension in ['jpg', 'jpeg', 'png'] else None,
                                    videos=[processed_file] if file_extension in ['mp4', 'mov', 'avi'] else None
                                )
                                
                                # Add web search if needed
                                if "additional information" in response.content.lower() or "cannot fully answer" in response.content.lower():
                                    with st.spinner("Searching for additional information..."):
                                        search_query = f"Additional information about {query}"
                                        search_results = perform_linkup_search(search_query, depth="deep")
                                        
                                        # Create supplementary prompt with search results
                                        supplementary_prompt = f"""
                                        CHAT HISTORY:
                                        {chat_context}
                                        
                                        Based on your initial analysis of the {'image' if file_extension in ['jpg', 'jpeg', 'png'] else 'video'},
                                        and the following additional information from web search:
                                        
                                        {search_results}
                                        
                                        Provide an updated response to the user query: {query}
                                        
                                        Consider our previous conversation for context.
                                        Clearly indicate when you're using information from the web search versus your direct analysis of the media.
                                        """
                                        
                                        # Get updated response with search results
                                        updated_response = multimodal_Agent.run(
                                            supplementary_prompt,
                                            images=[processed_file] if file_extension in ['jpg', 'jpeg', 'png'] else None,
                                            videos=[processed_file] if file_extension in ['mp4', 'mov', 'avi'] else None
                                        )
                                        
                                        answer = updated_response.content
                                else:
                                    answer = response.content
                        else:
                            # This is a general question not related to the media, use web search directly
                            with st.spinner("Searching for information to answer your query..."):
                                search_results = perform_linkup_search(query, depth="deep")
                                
                                # Create prompt for general question with search results - FIXED VERSION
                                general_prompt = f"""
                                CHAT HISTORY:
                                {chat_context}
                                
                                USER QUERY: {query}
                                
                                Based on the following search results and our previous conversation, please provide a helpful response:
                                
                                {search_results}
                                
                                FORMAT YOUR RESPONSE:
                                - Provide a direct, factual answer to the query
                                - Be concise but comprehensive
                                - Consider our previous conversation for context
                                """
                                
                                # Get response for general question - IMPORTANT: Use chat_Agent instead of multimodal_Agent
                                general_response = chat_Agent.run(general_prompt)
                                answer = general_response.content
            
        except Exception as error:
            answer = f"An error occurred: {error}"
        
        st.markdown(answer)
        st.session_state.chat_history.append({"role": "assistant", "content": answer})
        st.session_state.chat_memory.append({"role": "assistant", "content": answer})

# Customize text area height
st.markdown(
    """
    <style>
    .stTextArea textarea {
        height: 100px;
    }
    </style>
    """,
    unsafe_allow_html=True
)
